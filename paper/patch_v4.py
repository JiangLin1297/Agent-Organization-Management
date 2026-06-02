"""
Patch V4 papers: three academic rigor improvements.
1. Rewrite contribution claims (theory-inspired system engineering)
2. Add mechanism proof framing to §3.4
3. Add measurement scheme for Control-Emergence Balance Conjecture
"""

import os, sys
from docx import Document
from docx.shared import Pt


def get_style(doc, name):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles['Normal']


def replace_paragraph_text(para, new_text):
    """Replace text of a paragraph, preserving the first run's formatting."""
    if para.runs:
        fmt = para.runs[0].font
        size = fmt.size
        bold = fmt.bold
        italic = fmt.italic
        name = fmt.name
    else:
        size = bold = italic = name = None

    para.clear()
    run = para.add_run(new_text)
    if size:
        run.font.size = size
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if name:
        run.font.name = name


def insert_paragraph_after(doc, ref_para, text, style_name='Normal'):
    """Insert a new paragraph after ref_para."""
    new_para = doc.add_paragraph()
    ref_para._element.addnext(new_para._element)
    new_para.text = text
    new_para.style = get_style(doc, style_name)
    return new_para


def find_para_by_text(doc, text_prefix, style=None):
    """Find paragraph index by text prefix."""
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(text_prefix):
            if style is None or p.style.name == style:
                return i
    return -1


def patch_v4_en(path):
    """Apply all three patches to English V4."""
    doc = Document(path)

    # =====================================================================
    # WORK 1: Rewrite contribution claims
    # =====================================================================

    # --- Abstract: paragraph [4] ---
    abstract_idx = find_para_by_text(doc, 'This paper presents a three-group controlled experiment')
    if abstract_idx >= 0:
        replace_paragraph_text(doc.paragraphs[abstract_idx],
            'This paper presents a three-group controlled experiment that directly tests whether '
            'organizational management theory improves multi-agent system performance. We frame '
            'this work as theory-inspired system engineering: systematically mapping classical '
            'management theories to multi-agent collaboration protocols, and testing the boundary '
            'conditions of their applicability in LLM-based agent systems through controlled experiments.')

    # --- Introduction: paragraph [6] ---
    intro_para_idx = find_para_by_text(doc, 'Multi-agent systems (MAS) have become a dominant paradigm')
    if intro_para_idx >= 0:
        replace_paragraph_text(doc.paragraphs[intro_para_idx],
            'Multi-agent systems (MAS) have become a dominant paradigm for tackling complex, '
            'decomposed tasks. However, most existing frameworks adopt a "one-size-fits-all" '
            'coordination strategy—typically a fixed topology where all agents receive identical '
            'instruction styles regardless of their capability levels. This paper conducts '
            'theory-inspired system engineering: we systematically map Hersey and Blanchard\'s '
            'Situational Leadership Theory (1969) and related management theories to multi-agent '
            'system design, proposing AOM-DT (Dynamic Topology) as an adaptive alternative, and '
            'test the boundary conditions of these theories when applied to non-human agents.')

    # --- Conclusion: paragraph [133] ---
    concl_idx = find_para_by_text(doc, 'This paper has progressed from theoretical framework')
    if concl_idx >= 0:
        replace_paragraph_text(doc.paragraphs[concl_idx],
            'This paper has progressed from theoretical framework (V1) to simulated validation (V2) '
            'to real-world efficiency testing (V3) to controlled organizational comparison (V4). '
            'Each version answers a question that the previous version left open. Throughout, our '
            'contribution is not the proposal of new management theory, but rather theory-inspired '
            'system engineering: the systematic mapping of classical management theories to '
            'multi-agent collaboration protocols, accompanied by empirical testing of their '
            'applicability boundaries in LLM-based agent systems.')

    # --- After §3.7, before §4, add theoretical positioning paragraph ---
    v3v4_idx = find_para_by_text(doc, 'Together, V3 and V4 establish a complete picture')
    if v3v4_idx >= 0:
        insert_paragraph_after(doc, doc.paragraphs[v3v4_idx],
            'A note on theoretical contribution: this paper does not claim to have generated new '
            'management theory. Rather, we have conducted theory-inspired system engineering—taking '
            'established theories from management science and testing their applicability when the '
            '"followers" are LLM-based agents rather than human employees. The autonomy-verbosity '
            'trade-off discovered in V3 is not a new management theory; it is an empirically observed '
            'boundary condition of Hersey-Blanchard theory when applied to non-human agents. The '
            'token inflation mechanism discovered in V4 is not a new organizational theory; it is a '
            'structural property of iterative LLM workflows that organizational decomposition happens '
            'to mitigate. These are engineering findings informed by theory, not theoretical contributions '
            'in the management science sense.')

    # =====================================================================
    # WORK 2: Add mechanism proof framing to §3.4
    # =====================================================================

    # Find the §3.4 heading
    s34_idx = find_para_by_text(doc, '3.4 Token Inflation: The Structural Flaw')
    if s34_idx >= 0:
        # Insert mechanism proof framing after the existing §3.4 content
        # Find the last paragraph of §3.4 (before §3.5)
        s35_idx = find_para_by_text(doc, '3.5 The Testing Role')
        if s35_idx > 0:
            last_s34 = doc.paragraphs[s35_idx - 1]
            # Find the actual last non-empty paragraph before §3.5
            for j in range(s35_idx - 1, s34_idx, -1):
                if doc.paragraphs[j].text.strip():
                    last_s34 = doc.paragraphs[j]
                    break

            insert_paragraph_after(doc, last_s34,
                'The V4 experiment\'s core contribution is not which group produced better code, but '
                'the identification of a structural mechanism: the token inflation positive feedback '
                'loop inherent in single-agent iterative architectures. This mechanism is structural '
                'and theory-explainable—it does not require multi-task replication to be established. '
                'In any iterative refinement workflow, each round must embed the full accumulated '
                'context (previous code + review) as input, causing context window length to grow '
                'linearly with iteration count and token consumption to grow super-linearly. '
                'Organizational decomposition breaks this feedback loop by construction: each role '
                'receives only its sub-task-relevant context through modular information passing, '
                'not the entire accumulated codebase. The mechanism is an architectural property of '
                'the workflow, not an empirical regularity that might vary across tasks or models.')

    # =====================================================================
    # WORK 3: Add measurement scheme
    # =====================================================================

    # First, check if §2.4 exists. If not, we need to add it before §2.5.
    # In V4, §2 goes directly to §2.5. We need to add §2.4.4 before §2.5.
    s25_idx = find_para_by_text(doc, '2.5 Theoretical Solutions')
    if s25_idx > 0:
        # Find the last paragraph before §2.5
        insert_point = doc.paragraphs[s25_idx - 1]
        for j in range(s25_idx - 1, 8, -1):  # after §2 heading
            if doc.paragraphs[j].text.strip():
                insert_point = doc.paragraphs[j]
                break

        # Insert the measurement scheme content
        # We'll insert in reverse order since each insert goes after the reference
        measurement_content = [
            ('2.4 The Control-Emergence Balance: A Measurement Proposal', 'Heading 2'),
            ('The Control-Emergence Balance Conjecture (first proposed in V2) hypothesizes that '
             'multi-agent system performance follows an inverted-U relationship with control intensity: '
             'too little control leads to chaotic emergence; too much control suppresses beneficial '
             'emergent behaviors such as spontaneous role adaptation and creative task decomposition. '
             'The optimal operating point lies at an intermediate level of control that preserves '
             'structured coordination while allowing adaptive emergence.', 'Normal'),
            ('This conjecture has not been empirically validated across V1–V4. However, we can '
             'propose operationalized measurement methods for its two core variables, providing a '
             'concrete path for future verification.', 'Normal'),
            ('2.4.4 Measurement Proposal', 'Heading 3'),
            ('Control Intensity (CI): Operationalized as the ratio of management-protocol-constrained '
             'decision steps to total decision steps during task execution. In an agent system, this '
             'can be quantified by counting the proportion of steps where the agent follows predefined '
             'collaboration rules (reporting, approval, task division) versus steps where the agent '
             'acts autonomously. Formally:', 'Normal'),
            ('CI = N_constrained / N_total   where N_constrained = steps following predefined protocol, '
             'N_total = total decision steps   ... (4)', 'Normal'),
            ('A CI of 1.0 means full bureaucratic control (every step is protocol-governed); '
             'a CI of 0.0 means complete anarchy (no protocol约束). The conjecture predicts that '
             'system performance is maximized at some intermediate CI* ∈ (0, 1).', 'Normal'),
            ('Emergence Degree (ED): Operationalized as the proportion of agent output containing '
             'collaboration patterns that were not predefined in the management protocol. In an agent '
             'system, this can be quantified by detecting whether communication patterns deviate from '
             'the preset communication topology—for example, direct communication links between agents '
             'that were not in the preset hierarchy, or spontaneous role specialization not prescribed '
             'by the protocol. Formally:', 'Normal'),
            ('ED = N_novel / N_total   where N_novel = communication events outside preset topology, '
             'N_total = total communication events   ... (5)', 'Normal'),
            ('An ED of 0.0 means pure protocol compliance (no emergent behavior); an ED of 1.0 means '
             'complete protocol violation. The conjecture predicts that beneficial emergence clusters '
             'at intermediate ED values—enough to enable creative problem-solving, not enough to '
             'disrupt coordination.', 'Normal'),
            ('These measurement methods are not yet implemented. They are proposed as operationalizable '
             'paths for future empirical verification of the Control-Emergence Balance Conjecture. '
             'A V5 experiment could manipulate CI (by varying the strictness of collaboration protocols) '
             'and measure both ED and task performance, testing whether the predicted inverted-U '
             'relationship holds.', 'Normal'),
        ]

        # Insert in order (each after the previous)
        current_ref = insert_point
        for text, style in measurement_content:
            current_ref = insert_paragraph_after(doc, current_ref, text, style)

    doc.save(path)
    print(f'Patched EN: {path}')


def patch_v4_cn(path):
    """Apply all three patches to Chinese V4."""
    doc = Document(path)

    # =====================================================================
    # WORK 1: Rewrite contribution claims (CN)
    # =====================================================================

    # Abstract
    abstract_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == 'Heading 1' and '摘要' in p.text:
            abstract_idx = i
            break
    if abstract_idx >= 0:
        # The abstract is the next paragraph after the heading
        for j in range(abstract_idx + 1, abstract_idx + 5):
            if doc.paragraphs[j].text.strip() and doc.paragraphs[j].style.name == 'Normal':
                old = doc.paragraphs[j].text
                if '三组对照实验' in old and '直接检验' in old:
                    replace_paragraph_text(doc.paragraphs[j],
                        '本文报告了一项三组对照实验，直接检验组织管理学理论是否能提升多智能体系统的表现。'
                        '我们将本文定位为理论启发的系统工程（theory-inspired system engineering）：'
                        '将经典管理理论系统性地映射为多Agent协作协议，并通过受控实验检验这些理论在LLM Agent情境下的适用性边界。')
                break

    # Introduction
    intro_idx = find_para_by_text(doc, '当人工智能代理')
    if intro_idx >= 0:
        replace_paragraph_text(doc.paragraphs[intro_idx],
            '当人工智能代理（AI Agent）从孤立的工具演化为"数字员工"时，管理它们的最优解不再是编程，'
            '而是管理学本身。当前主流的多智能体框架（如AutoGen、CrewAI、LangGraph）普遍存在"管理真空"'
            '——协作流程以静态拓扑硬编码，缺乏权变响应能力。本文进行了一项理论启发的系统工程：'
            '将赫塞-布兰查德的情境领导理论及相关管理理论系统性地映射为多智能体协作协议，'
            '并通过受控实验检验这些理论在迁移到非人类主体时的适用性边界。')

    # Conclusion
    concl_idx = find_para_by_text(doc, '本文从理论框架')
    if concl_idx >= 0:
        replace_paragraph_text(doc.paragraphs[concl_idx],
            '本文从理论框架（V1）到模拟验证（V2）到效率测试（V3）到组织对照实验（V4），'
            '逐步回答了每个版本遗留的问题。贯穿始终，我们的贡献不是提出新的管理理论，'
            '而是理论启发的系统工程：将经典管理理论系统性地映射为多Agent协作协议，'
            '并通过实证实验检验这些理论在LLM Agent系统中的适用性边界。')

    # After §3.7 add theoretical positioning
    v3v4_idx = find_para_by_text(doc, '二者共同确立了完整的图景')
    if v3v4_idx >= 0:
        insert_paragraph_after(doc, doc.paragraphs[v3v4_idx],
            '关于理论贡献的说明：本文不声称提出了新的管理理论。我们进行的是理论启发的系统工程'
            '——从管理学中提取成熟理论，检验其在"追随者"为LLM Agent而非人类员工时的适用性。'
            'V3发现的"自主性-冗长度权衡"不是新的管理理论，而是赫塞-布兰查德理论迁移到非人类主体时'
            '的一个此前未被文献记录的适用性边界条件。V4发现的"Token膨胀机制"不是新的组织理论，'
            '而是迭代式LLM工作流的结构性特性——组织分工恰好能规避它。'
            '这些是由理论启发的工程发现，而非管理学意义上的理论贡献。')

    # =====================================================================
    # WORK 2: Add mechanism proof framing (CN)
    # =====================================================================

    s34_idx = find_para_by_text(doc, '3.4 Token膨胀的根因分析')
    if s34_idx >= 0:
        s35_idx = find_para_by_text(doc, '3.5 测试环节发现')
        if s35_idx > 0:
            last_s34 = doc.paragraphs[s35_idx - 1]
            for j in range(s35_idx - 1, s34_idx, -1):
                if doc.paragraphs[j].text.strip():
                    last_s34 = doc.paragraphs[j]
                    break

            insert_paragraph_after(doc, last_s34,
                'V4实验的核心贡献不在于"哪个组质量更好"，而在于揭示了一个结构性机制——'
                '单Agent迭代架构存在Token膨胀的正反馈循环：每次迭代需要将前一轮完整代码作为输入，'
                '导致上下文窗口随迭代次数线性增长，Token消耗呈超线性增长。'
                '而组织分工架构通过模块化信息传递（每个角色只接收其子任务相关的上下文），'
                '从根本上规避了这个膨胀机制。这一机制是结构性的、理论可解释的——'
                '它不依赖于特定任务或特定模型，而是迭代架构的固有属性。')

    # =====================================================================
    # WORK 3: Add measurement scheme (CN)
    # =====================================================================

    s25_idx = find_para_by_text(doc, '2.5 自主性-冗长度权衡')
    if s25_idx > 0:
        insert_point = doc.paragraphs[s25_idx - 1]
        for j in range(s25_idx - 1, 8, -1):
            if doc.paragraphs[j].text.strip():
                insert_point = doc.paragraphs[j]
                break

        measurement_content = [
            ('2.4 控制-涌现平衡：一个测量方案', 'Heading 2'),
            ('控制-涌现平衡猜想（在V2中首次提出）假设，多智能体系统性能与控制强度呈倒U型关系：'
             '过低的控制导致混沌涌现，过高的控制压制有益的涌现行为（如自发角色适应和创造性任务分解）。'
             '最优工作点位于中等控制水平——既保持结构化协调，又允许自适应涌现。', 'Normal'),
            ('该猜想在V1-V4中尚未获得实证验证。但我们可以为其两个核心变量提出可操作的测量方式，'
             '为未来的验证提供具体路径。', 'Normal'),
            ('2.4.4 测量方案', 'Heading 3'),
            ('控制强度（Control Intensity, CI）：操作化为"受管理协议约束的决策步骤数 / 总决策步骤数"。'
             '在Agent系统中，可以通过统计Agent在任务执行过程中遵循预定义协作规则'
             '（汇报、审批、分工）的步骤比例来量化。形式化表示为：', 'Normal'),
            ('CI = N_约束 / N_总   其中 N_约束 = 遵循预定义协议的步骤数，'
             'N_总 = 总决策步骤数   …… (4)', 'Normal'),
            ('CI=1.0表示完全科层制控制（每一步都受协议约束），CI=0.0表示完全无政府状态。'
             '猜想预测系统性能在某个中间CI*∈(0,1)处最大化。', 'Normal'),
            ('涌现度（Emergence Degree, ED）：操作化为"Agent输出中超出预定义协议的新协作模式的比例"。'
             '在Agent系统中，可以通过检测Agent之间的通信模式是否偏离了预设的通信拓扑来量化'
             '——如出现了预设之外的直接通信链路，或自发形成了预设之外的角色分工。形式化表示为：', 'Normal'),
            ('ED = N_新颖 / N_总   其中 N_新颖 = 预设拓扑之外的通信事件数，'
             'N_总 = 总通信事件数   …… (5)', 'Normal'),
            ('ED=0.0表示纯粹的协议遵从（无涌现行为），ED=1.0表示完全的协议违反。'
             '猜想预测有益涌现聚集在中等ED值——足以支持创造性问题解决，又不至于破坏协调。', 'Normal'),
            ('这些测量方式目前尚未实现。它们是为未来实证验证控制-涌现平衡猜想而提出的可操作路径。'
             'V5实验可以通过操纵CI（改变协作协议的严格程度）并同时测量ED和任务表现，'
             '检验预测的倒U型关系是否成立。', 'Normal'),
        ]

        current_ref = insert_point
        for text, style in measurement_content:
            current_ref = insert_paragraph_after(doc, current_ref, text, style)

    doc.save(path)
    print(f'Patched CN: {path}')


if __name__ == '__main__':
    base = os.path.dirname(os.path.abspath(__file__))
    en = os.path.join(base, 'agent_organizational_management_v4.docx')
    cn = os.path.join(base, 'agent_organizational_management_v4_cn.docx')

    patch_v4_en(en)
    patch_v4_cn(cn)
    print('Done. All three patches applied to both papers.')
