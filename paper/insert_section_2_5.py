"""
Insert Section 2.5 into V3 papers (English and Chinese).
Theoretical solutions to the autonomy-verbosity tradeoff.
"""

from docx import Document
from docx.shared import Pt
import re


def find_paragraph_index(doc, text_prefix, style_name=None):
    """Find paragraph index by text prefix and optional style."""
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith(text_prefix):
            if style_name is None or para.style.name == style_name:
                return i
    return -1


def get_style(doc, style_name):
    try:
        return doc.styles[style_name]
    except KeyError:
        return doc.styles['Normal']


def insert_paragraph_before_element(doc, ref_para, text, style_name='Normal'):
    """Insert a paragraph before a reference paragraph."""
    new_para = doc.add_paragraph()
    ref_para._element.addprevious(new_para._element)
    new_para.text = text
    new_para.style = get_style(doc, style_name)
    return new_para


def build_en_section_content():
    """Build English section 2.5 content as list of (text, style) tuples."""
    content = []

    content.append((
        "2.5 Theoretical Solutions to the Autonomy-Verbosity Trade-off: From Management Principles to Prompt Engineering",
        "Heading 2"
    ))

    content.append((
        "Section 5.2 identifies the autonomy-verbosity trade-off as the central efficiency bottleneck of AOM-DT: "
        "high-autonomy instruction styles (S3/S4) elicit more verbose, exploratory outputs from capable LLM-based agents, "
        "increasing token consumption without proportional quality gains. This section proposes a systematic solution "
        "grounded in management theory. We select four classical theories—not arbitrarily, but because each addresses "
        "a distinct facet of the verbosity problem—and engineer them into concrete prompt design principles.",
        "Normal"
    ))

    content.append((
        "The core insight is this: in human management, the problem of subordinate verbosity is not new. "
        "Managers have faced it for over a century, and management science has developed well-tested solutions. "
        "The autonomy-verbosity trade-off is, at its root, an information flow control problem—and management "
        "theory offers a rich toolkit for governing information flow in hierarchical organizations.",
        "Normal"
    ))

    # 2.5.2
    content.append(("2.5.2 Theory Selection and Rationale", "Heading 3"))

    content.append((
        "We select four management theories based on their direct relevance to the verbosity problem:",
        "Normal"
    ))

    content.append((
        "1. Toyota’s A3 Report (Shook, 2008): A structured one-page reporting framework that constrains "
        "output format and length. Selected because it directly solves the unstructured verbosity problem—"
        "when an agent receives a high-autonomy prompt, the absence of structural constraints allows unconstrained "
        "generation. A3 provides the structural scaffold.",
        "Normal"
    ))

    content.append((
        "2. Management by Exception (Drucker, 1954): A control principle where managers intervene only when "
        "performance deviates from standards. Selected because it addresses the report-everything problem—"
        "current LLMs, when given autonomy, default to exhaustive coverage rather than selective reporting.",
        "Normal"
    ))

    content.append((
        "3. Simon’s Bounded Rationality and Satisficing (Simon, 1955): The principle that decision-makers "
        "should seek good enough solutions rather than optimal ones, given cognitive constraints. Selected because "
        "it addresses the diminishing returns problem—LLMs continue elaborating well past the point of marginal value.",
        "Normal"
    ))

    content.append((
        "4. Mintzberg’s Informational Roles (Mintzberg, 1973): The theory that managerial information processing "
        "involves three roles—monitor, disseminator, and spokesperson—each with distinct filtering criteria. "
        "Selected because it addresses the information relevance problem—LLMs lack the ability to distinguish "
        "high-value from low-value information in their outputs.",
        "Normal"
    ))

    # 2.5.3
    content.append(("2.5.3 Theory-to-Prompt Engineering Mappings", "Heading 3"))

    # Mapping 1: A3
    content.append(("Mapping 1: A3 Report → Structured Output Protocol", "Heading 3"))

    content.append((
        "The Toyota A3 Report constrains problem-solving communication to a single A3-sized page with fixed sections: "
        "background, current state, analysis, target state, implementation plan, and follow-up (Shook, 2008). "
        "The constraint is not merely physical—it enforces a discipline of conciseness: every word must earn its place.",
        "Normal"
    ))

    content.append((
        "Engineering mapping: Transform open-ended S4 prompts into structured-output prompts with fixed sections "
        "and explicit length budgets.",
        "Normal"
    ))

    content.append(("Unconstrained S4 (current):", "Normal"))
    content.append((
        "“You are a senior agent. Handle this task independently. Use your expertise to deliver the best solution.”",
        "Normal"
    ))

    content.append(("A3-constrained S4 (proposed):", "Normal"))
    content.append((
        "“You are a senior agent. Handle this task independently. Structure your response as follows:\n"
        "[1] Problem Statement (≤50 words)\n"
        "[2] Key Analysis (≤200 words, max 3 points)\n"
        "[3] Recommended Solution (≤150 words)\n"
        "[4] Implementation Steps (≤100 words, max 5 steps)\n"
        "Total output must not exceed 800 words. Eliminate redundancy and filler phrases.”",
        "Normal"
    ))

    content.append((
        "Design Principle 1 (Structural Constraint): High-autonomy prompts must include a structured output template "
        "with explicit section word counts. The total token budget should be calibrated to the task’s information "
        "density: low-uncertainty tasks ≤ 500 words, medium ≤ 800, high ≤ 1200.",
        "Normal"
    ))

    # Mapping 2: MBE
    content.append(("Mapping 2: Management by Exception → Deviation-Based Reporting", "Heading 3"))

    content.append((
        "Management by Exception (MBE) stipulates that managers should focus attention on significant deviations from "
        "planned performance, rather than reviewing all routine operations (Drucker, 1954). In production management, "
        "this means workers report only when outcomes fall outside predefined tolerance bands.",
        "Normal"
    ))

    content.append((
        "Engineering mapping: Define normal and exceptional conditions in the prompt. "
        "Instruct the agent to elaborate only on deviations.",
        "Normal"
    ))

    content.append(("Unconstrained S4:", "Normal"))
    content.append((
        "“Analyze this system architecture and provide a comprehensive review.”",
        "Normal"
    ))

    content.append(("MBE-constrained S4:", "Normal"))
    content.append((
        "“Review this system architecture. For components that meet the following standards, state only: "
        "‘Component X: Meets standard.’ Elaborate ONLY on deviations:\n"
        "- Response time < 200ms\n- Availability > 99.9%\n- Error rate < 0.1%\n"
        "For each deviation, provide: (1) the gap, (2) root cause, (3) recommended fix. Maximum 3 deviations.”",
        "Normal"
    ))

    content.append((
        "Design Principle 2 (Deviation Filtering): High-autonomy prompts must define explicit success criteria or "
        "normal baselines. The agent should be instructed to report concisely on conforming items and elaborate only "
        "on exceptions. This inverts the default LLM behavior of uniform-depth analysis.",
        "Normal"
    ))

    # Mapping 3: Simon
    content.append(("Mapping 3: Bounded Rationality → Satisficing Criteria", "Heading 3"))

    content.append((
        "Simon’s bounded rationality argues that rational decision-making is constrained by cognitive limitations, "
        "available information, and time. Decision-makers do not optimize—they satisfice, accepting the first "
        "solution that meets a minimum acceptable threshold (Simon, 1955). This is not a compromise; it is rational "
        "behavior under constraints.",
        "Normal"
    ))

    content.append((
        "Engineering mapping: Define explicit satisficing thresholds in the prompt. "
        "Instruct the agent to stop analysis once the threshold is met.",
        "Normal"
    ))

    content.append(("Unconstrained S3:", "Normal"))
    content.append((
        "“Participate in solving this problem. Explore multiple approaches and discuss their trade-offs in detail.”",
        "Normal"
    ))

    content.append(("Satisficing-constrained S3:", "Normal"))
    content.append((
        "“Participate in solving this problem. Evaluate candidate approaches against these criteria: "
        "(1) feasibility > 0.7, (2) cost < budget, (3) time-to-implement < 2 weeks. Select the FIRST approach "
        "that meets all three criteria. Present your selection with a 100-word justification. Do not continue "
        "evaluating after finding a satisfactory solution.”",
        "Normal"
    ))

    content.append((
        "Design Principle 3 (Satisficing Stopping Rule): High-autonomy prompts must include explicit stopping criteria. "
        "The agent should terminate analysis upon finding a solution that meets predefined thresholds, rather than "
        "exhaustively exploring the solution space. This directly counters the LLM tendency toward exhaustive elaboration.",
        "Normal"
    ))

    # Mapping 4: Mintzberg
    content.append(("Mapping 4: Informational Roles → Information Filtering Protocol", "Heading 3"))

    content.append((
        "Mintzberg identified three informational roles for managers: monitor (scanning for relevant information), "
        "disseminator (forwarding information to subordinates), and spokesperson (transmitting information to outsiders) "
        "(Mintzberg, 1973). Critically, the monitor role is not about collecting all information—it is about "
        "selective attention to strategically relevant signals.",
        "Normal"
    ))

    content.append((
        "Engineering mapping: Define an explicit information filter in the prompt that specifies what constitutes "
        "high-value vs. low-value output.",
        "Normal"
    ))

    content.append(("Unconstrained S4:", "Normal"))
    content.append((
        "“Design a database schema for this e-commerce platform. Provide a thorough analysis.”",
        "Normal"
    ))

    content.append(("Filter-constrained S4:", "Normal"))
    content.append((
        "“Design a database schema for this e-commerce platform. Your output must prioritize:\n"
        "HIGH VALUE: Schema design decisions that affect query performance (>10ms impact), data integrity "
        "constraints, and scalability bottlenecks.\n"
        "LOW VALUE (omit): Generic best practices, obvious design choices, and implementation details that "
        "don’t affect the architecture.\n"
        "Focus your analysis on the top 3 highest-impact decisions only.”",
        "Normal"
    ))

    content.append((
        "Design Principle 4 (Information Filtering): High-autonomy prompts must include an explicit value hierarchy "
        "that distinguishes high-value from low-value information. The agent should allocate output length proportional "
        "to information value, not uniformly across all aspects.",
        "Normal"
    ))

    # 2.5.4
    content.append(("2.5.4 Unified Constrained Delegation Prompt Template", "Heading 3"))

    content.append((
        "Integrating the four mappings, we propose a unified prompt template for constrained high-autonomy instructions:",
        "Normal"
    ))

    content.append(("Template: Constrained S4 (Delegating) Prompt", "Normal"))

    content.append((
        "“You are a senior agent with readiness level R4. Handle this task independently.\n\n"
        "OBJECTIVE: [specific, measurable goal]\n\n"
        "OUTPUT STRUCTURE (A3 Protocol):\n"
        "- Problem Statement: ≤ [N1] words\n"
        "- Key Analysis: ≤ [N2] words, max [K] points\n"
        "- Recommendation: ≤ [N3] words\n"
        "- Next Steps: ≤ [N4] words\n"
        "- Total: must not exceed [T] words\n\n"
        "DEVIATION FILTER (MBE):\n"
        "- Report routine items as one-line status.\n"
        "- Elaborate ONLY on items that deviate from: [criteria]\n"
        "- Max deviations to discuss: [D]\n\n"
        "STOPPING RULE (Satisficing):\n"
        "- Stop analysis upon finding a solution that meets: [thresholds]\n"
        "- Do not continue evaluating after finding a satisfactory option.\n\n"
        "VALUE HIERARCHY (Information Filter):\n"
        "- HIGH VALUE: [criteria]\n"
        "- LOW VALUE (omit): [criteria]\n"
        "- Allocate detail proportional to value.”",
        "Normal"
    ))

    # 2.5.5
    content.append(("2.5.5 Model-Agnostic Implementation", "Heading 3"))

    content.append((
        "The proposed design principles are model-agnostic—they rely on prompt-level constraints rather than "
        "model-specific features. Implementation requires no fine-tuning, no special API parameters, and no "
        "model-specific workarounds. The approach works with any instruction-following LLM (GPT-4, Claude, MiMo, "
        "Llama, etc.) because it leverages a universal property: LLMs respond to structural cues in prompts.",
        "Normal"
    ))

    content.append(("Implementation strategy:", "Normal"))

    content.append((
        "1. Prompt-layer integration: Embed the constrained template into the AOM-DT style assignment logic. "
        "When the readiness calculator assigns S3 or S4, the system automatically appends the appropriate constraints "
        "(A3 structure, MBE filter, satisficing threshold, information hierarchy) to the base instruction.",
        "Normal"
    ))

    content.append((
        "2. Token budget calibration: The word limits [N1]-[N4] and total [T] should be calibrated per task "
        "uncertainty level. Based on our V3 data, we propose:",
        "Normal"
    ))

    content.append(("   Low uncertainty (0.1): T = 500 words (≈ 650 tokens)", "Normal"))
    content.append(("   Medium uncertainty (0.5): T = 800 words (≈ 1040 tokens)", "Normal"))
    content.append(("   High uncertainty (0.9): T = 1200 words (≈ 1560 tokens)", "Normal"))

    content.append((
        "3. Deviation criteria specification: The MBE deviation criteria should be derived from the task specification. "
        "For each task, the coordinator agent (Agent D) generates task-specific success criteria during the coordination "
        "phase, which are then passed to workers as MBE baselines.",
        "Normal"
    ))

    content.append((
        "4. Graceful degradation: If the LLM exceeds the word budget, the system should truncate at the budget "
        "boundary and append a note: [Output truncated at word budget. Key remaining points: ...] This ensures "
        "the constraint is enforced even when the LLM does not perfectly comply.",
        "Normal"
    ))

    # 2.5.6
    content.append(("2.5.6 Theoretical Synthesis: From Verbosity Control to Optimal Information Flow", "Heading 3"))

    content.append((
        "The four management theories converge on a unified principle: optimal information flow requires explicit "
        "constraints on what, when, and how much to communicate. In human organizations, these constraints are "
        "internalized through training, culture, and professional norms. LLMs lack this internalization—their "
        "training (pre-training on internet text) optimizes for fluency and comprehensiveness, not for concise, "
        "targeted communication under organizational constraints.",
        "Normal"
    ))

    content.append((
        "The proposed solution externalizes these constraints into the prompt itself. This is analogous to how "
        "organizations use standardized reporting templates, communication protocols, and performance dashboards to "
        "govern information flow—tools that Fayol (1916) would recognize as extensions of the unity of direction principle.",
        "Normal"
    ))

    content.append((
        "Formally, let V(S, L) denote the verbosity of LLM L under instruction style S. The autonomy-verbosity "
        "trade-off observed in V3 can be expressed as:",
        "Normal"
    ))

    content.append((
        "V(S4, L) > V(S1, L)  for current LLMs   ... (1)",
        "Normal"
    ))

    content.append((
        "The constrained delegation approach aims to achieve:",
        "Normal"
    ))

    content.append((
        "V(S4_constrained, L) ≈ V(S1, L)  while preserving  Quality(S4, L) ≥ Quality(S1, L)   ... (2)",
        "Normal"
    ))

    content.append((
        "Equation (2) states the design goal: constrained high-autonomy prompts should match the token efficiency "
        "of low-autonomy prompts while preserving the quality advantages of autonomous execution. This is achievable "
        "because the verbosity in S4 is not caused by the task’s inherent complexity—it is an artifact of "
        "unconstrained generation. The task’s information content is fixed; only the output format changes.",
        "Normal"
    ))

    content.append((
        "We hypothesize that the constrained approach will reduce AOM-DT’s token overhead from +7.0% to within "
        "+2% of ST, while maintaining the 100% success rate and ~3.9/5 quality observed in V3. "
        "This hypothesis is testable in a V4 experiment.",
        "Normal"
    ))

    return content


def build_cn_section_content():
    """Build Chinese section 2.5 content as list of (text, style) tuples."""
    content = []

    content.append((
        "2.5 自主性-冗长度权衡的理论解：从管理原则到Prompt工程",
        "Heading 2"
    ))

    content.append((
        "5.2节将自主性-冗长度权衡识别为AOM-DT的核心效率瓶颈："
        "高自主性指令风格（S3/S4）会激发有能力的LLM智能体产生更冗长、更具探索性的输出，"
        "增加Token消耗却未带来等比例的质量提升。本节从管理学理论出发，提出系统性的解决方案。"
        "我们选择四个经典理论——不是任意选择，而是因为每个理论分别针对冗长度问题的不同侧面"
        "——并将其工程化为具体的Prompt设计原则。",
        "Normal"
    ))

    content.append((
        "核心洞见是：在人类管理中，“下属冗长”并不是新问题。"
        "管理者面对这个问题已逾百年，管理学已经发展出了经过充分检验的解决方案。"
        "自主性-冗长度权衡本质上是一个信息流控制问题——而管理学为治理层级组织中的信息流提供了丰富的工具箱。",
        "Normal"
    ))

    # 2.5.2
    content.append(("2.5.2 理论选择与依据", "Heading 3"))

    content.append((
        "我们基于与冗长度问题的直接相关性，选择以下四个管理学理论：",
        "Normal"
    ))

    content.append((
        "1. 丰田A3报告（Shook, 2008）：一种结构化的单页报告框架，"
        "约束输出格式与长度。选择理由：它直接解决“非结构化冗长”问题——"
        "当智能体收到高自主性Prompt时，缺乏结构约束会导致不受控的生成。A3提供了结构化骨架。",
        "Normal"
    ))

    content.append((
        "2. 例外管理（Drucker, 1954）：管理者仅在绩效偏离标准时才介入的控制原则。"
        "选择理由：它解决“事无巨细”问题——"
        "当前LLM在获得自主权时，默认倾向于全面覆盖而非选择性报告。",
        "Normal"
    ))

    content.append((
        "3. 西蒙的有限理性与满意化（Simon, 1955）："
        "决策者应在认知约束下寻找“足够好”的方案而非最优方案的原则。"
        "选择理由：它解决“边际收益递减”问题——"
        "LLM会持续阐述，远超边际价值消失的临界点。",
        "Normal"
    ))

    content.append((
        "4. 明茨伯格的信息角色理论（Mintzberg, 1973）："
        "管理者的信息处理涉及三个角色——监听者、传播者和发言人，"
        "每个角色有不同的过滤标准。选择理由：它解决“信息相关性”问题——"
        "LLM缺乏区分输出中高价值与低价值信息的能力。",
        "Normal"
    ))

    # 2.5.3
    content.append(("2.5.3 理论到Prompt工程的映射", "Heading 3"))

    # Mapping 1: A3
    content.append(("映射一：A3报告 → 结构化输出协议", "Heading 3"))

    content.append((
        "丰田A3报告将问题解决的沟通约束在一张A3纸上，"
        "包含固定板块：背景、现状、分析、目标状态、实施计划和后续跟踪（Shook, 2008）。"
        "这一约束不仅是物理层面的——它强制执行一种精练的表达纪律：每一个字都必须证明其存在的价值。",
        "Normal"
    ))

    content.append((
        "工程映射：将开放式S4 Prompt转换为带有固定板块和显式长度预算的结构化输出Prompt。",
        "Normal"
    ))

    content.append(("无约束S4（当前）：", "Normal"))
    content.append((
        "“你是一名资深智能体。独立处理此任务。运用你的专业能力提供最佳方案。”",
        "Normal"
    ))

    content.append(("A3约束S4（改进）：", "Normal"))
    content.append((
        "“你是一名资深智能体。独立处理此任务。按以下结构组织输出：\n"
        "[1] 问题陈述（≤50字）\n"
        "[2] 核心分析（≤200字，最多3个要点）\n"
        "[3] 推荐方案（≤150字）\n"
        "[4] 实施步骤（≤100字，最多5步）\n"
        "总输出不超过800字。消除冗余和填充性表述。”",
        "Normal"
    ))

    content.append((
        "设计原则一（结构约束）：高自主性Prompt必须包含带有显式板块字数预算的结构化输出模板。"
        "总Token预算应根据任务信息密度校准：低不确定性任务≤500字，中等≤800字，高≤1200字。",
        "Normal"
    ))

    # Mapping 2: MBE
    content.append(("映射二：例外管理 → 偏差报告机制", "Heading 3"))

    content.append((
        "例外管理（MBE）规定，管理者应将注意力集中在与计划绩效的显著偏差上，"
        "而非审查所有常规运营（Drucker, 1954）。在生产管理中，"
        "这意味着工人仅在结果超出预定义容差范围时才进行报告。",
        "Normal"
    ))

    content.append((
        "工程映射：在Prompt中定义“正常”和“异常”条件，"
        "指示智能体仅在偏差出现时才详细阐述。",
        "Normal"
    ))

    content.append(("无约束S4：", "Normal"))
    content.append((
        "“分析此系统架构并提供全面评审。”",
        "Normal"
    ))

    content.append(("MBE约束S4：", "Normal"))
    content.append((
        "“评审此系统架构。对于满足以下标准的组件，仅陈述：‘组件X：符合标准。’仅对偏差进行详细说明：\n"
        "- 响应时间 < 200ms\n- 可用性 > 99.9%\n- 错误率 < 0.1%\n"
        "对每个偏差提供：（1）差距（2）根因（3）建议修复。最多讨论3个偏差。”",
        "Normal"
    ))

    content.append((
        "设计原则二（偏差过滤）：高自主性Prompt必须定义显式的成功标准或“正常”基线。"
        "智能体应被指示对符合标准的项目进行简洁报告，仅对例外进行详细阐述。"
        "这逆转了LLM默认的等深度分析行为。",
        "Normal"
    ))

    # Mapping 3: Simon
    content.append(("映射三：有限理性 → 满意化停止规则", "Heading 3"))

    content.append((
        "西蒙的有限理性理论认为，理性决策受到认知限制、"
        "可用信息和时间的约束。决策者不进行最优化——他们“满意化”，"
        "接受第一个满足最低可接受阈值的方案（Simon, 1955）。"
        "这不是妥协，而是在约束条件下的理性行为。",
        "Normal"
    ))

    content.append((
        "工程映射：在Prompt中定义显式的满意化阈值，"
        "指示智能体在阈值满足时停止分析。",
        "Normal"
    ))

    content.append(("无约束S3：", "Normal"))
    content.append((
        "“参与解决此问题。探索多种方法并详细讨论其权衡。”",
        "Normal"
    ))

    content.append(("满意化约束S3：", "Normal"))
    content.append((
        "“参与解决此问题。根据以下标准评估候选方法："
        "（1）可行性 > 0.7（2）成本 < 预算（3）实施时间 < 2周。"
        "选择第一个满足所有三个标准的方法。用100字说明选择理由。"
        "找到满意方案后不要继续评估。”",
        "Normal"
    ))

    content.append((
        "设计原则三（满意化停止规则）：高自主性Prompt必须包含显式停止标准。"
        "智能体应在找到满足预定义阈值的方案时终止分析，"
        "而非穷尽地探索解空间。这直接对抗LLM倾向于穷尽阐述的行为。",
        "Normal"
    ))

    # Mapping 4: Mintzberg
    content.append(("映射四：信息角色 → 信息过滤协议", "Heading 3"))

    content.append((
        "明茨伯格识别了管理者的三个信息角色："
        "监听者（扫描相关信息）、传播者（向下属转发信息）和发言人（向外部传递信息）"
        "（Mintzberg, 1973）。关键在于，监听角色并非收集所有信息——"
        "而是对战略相关信号的选择性注意。",
        "Normal"
    ))

    content.append((
        "工程映射：在Prompt中定义显式的“信息过滤器”，"
        "区分高价值与低价值输出。",
        "Normal"
    ))

    content.append(("无约束S4：", "Normal"))
    content.append((
        "“为此电商平台设计数据库Schema。进行透彻的分析。”",
        "Normal"
    ))

    content.append(("过滤约束S4：", "Normal"))
    content.append((
        "“为此电商平台设计数据库Schema。输出必须优先考虑：\n"
        "高价值：影响查询性能（>10ms影响）的Schema设计决策、数据完整性约束、可扩展性瓶颈。\n"
        "低价值（省略）：通用最佳实践、显而易见的设计选择、不影响架构的实现细节。\n"
        "仅聚焦于前3个最高影响的决策。”",
        "Normal"
    ))

    content.append((
        "设计原则四（信息过滤）：高自主性Prompt必须包含显式的价值层次结构，"
        "区分高价值与低价值信息。智能体应按信息价值而非均匀地分配输出篇幅。",
        "Normal"
    ))

    # 2.5.4
    content.append(("2.5.4 统一约束授权Prompt模板", "Heading 3"))

    content.append((
        "综合四个映射，我们提出一个统一的高自主性指令Prompt模板：",
        "Normal"
    ))

    content.append(("模板：约束S4（授权式）Prompt", "Normal"))

    content.append((
        "“你是一名准备度为R4的资深智能体。独立处理此任务。\n\n"
        "目标（MBO）：[具体、可衡量的目标]\n\n"
        "输出结构（A3协议）：\n"
        "- 问题陈述：≤ [N1] 字\n"
        "- 核心分析：≤ [N2] 字，最多 [K] 个要点\n"
        "- 推荐方案：≤ [N3] 字\n"
        "- 下一步行动：≤ [N4] 字\n"
        "- 总计：不超过 [T] 字\n\n"
        "偏差过滤（MBE）：\n"
        "- 常规项目以一句话状态报告\n"
        "- 仅对偏离以下标准的项目详细说明：[标准]\n"
        "- 最多讨论偏差数：[D]\n\n"
        "停止规则（满意化）：\n"
        "- 找到满足以下条件的方案后停止分析：[阈值]\n"
        "- 找到满意方案后不要继续评估\n\n"
        "价值层次（信息过滤）：\n"
        "- 高价值：[标准]\n"
        "- 低价值（省略）：[标准]\n"
        "- 按价值比例分配篇幅。”",
        "Normal"
    ))

    # 2.5.5
    content.append(("2.5.5 模型无关的实现方案", "Heading 3"))

    content.append((
        "所提出的设计原则是模型无关的——它们依赖Prompt层面的约束而非模型特性。"
        "实现无需微调、无需特殊API参数、无需模型特定的变通方案。"
        "该方案适用于任何遵循指令的LLM（GPT-4、Claude、MiMo、Llama等），"
        "因为它利用了一个通用特性：LLM对Prompt中的结构化线索做出响应。",
        "Normal"
    ))

    content.append(("实现策略：", "Normal"))

    content.append((
        "1. Prompt层集成：将约束模板嵌入AOM-DT的风格分配逻辑中。"
        "当准备度计算器分配S3或S4时，系统自动将相应约束"
        "（A3结构、MBE过滤器、满意化阈值、信息层次）附加到基础指令中。",
        "Normal"
    ))

    content.append((
        "2. Token预算校准：字数限制[N1]-[N4]和总计[T]应按任务不确定性水平校准。"
        "基于我们的V3数据，建议如下：",
        "Normal"
    ))

    content.append(("   低不确定性（0.1）：T = 500字（≈650 Token）", "Normal"))
    content.append(("   中不确定性（0.5）：T = 800字（≈1040 Token）", "Normal"))
    content.append(("   高不确定性（0.9）：T = 1200字（≈1560 Token）", "Normal"))

    content.append((
        "3. 偏差标准规范：MBE偏差标准应从任务规范中导出。"
        "对于每个任务，协调者智能体（Agent D）在协调阶段生成任务特定的成功标准，"
        "然后作为MBE基线传递给工作者。",
        "Normal"
    ))

    content.append((
        "4. 优雅降级：如果LLM超出字数预算，系统应在预算边界处截断并附加说明："
        "“[输出在字数预算处截断。剩余关键要点：……]”"
        "这确保即使LLM不完全遵守，约束也能得到强制执行。",
        "Normal"
    ))

    # 2.5.6
    content.append(("2.5.6 理论综合：从冗长度控制到最优信息流", "Heading 3"))

    content.append((
        "四个管理学理论汇聚于一个统一原则："
        "最优信息流需要对沟通的内容、时机和数量进行显式约束。"
        "在人类组织中，这些约束通过培训、文化和职业规范被内化。"
        "LLM缺乏这种内化——它们的“训练”（在互联网文本上的预训练）"
        "优化的是流畅性和全面性，而非在组织约束下的精练、有针对性的沟通。",
        "Normal"
    ))

    content.append((
        "所提出的方案将这些约束外化到Prompt本身。"
        "这类似于组织如何使用标准化报告模板、沟通协议和绩效仪表板来治理信息流——"
        "这些工具是法约尔（Fayol, 1916）会认可的统一指挥原则的延伸。",
        "Normal"
    ))

    content.append((
        "形式化地，令V(S, L)表示LLM L在指令风格S下的冗长度。"
        "V3中观察到的自主性-冗长度权衡可表示为：",
        "Normal"
    ))

    content.append((
        "V(S4, L) > V(S1, L)，对当前LLM成立   …… (1)",
        "Normal"
    ))

    content.append(("约束授权方案旨在实现：", "Normal"))

    content.append((
        "V(S4_约束, L) ≈ V(S1, L)，同时保持 Quality(S4, L) ≥ Quality(S1, L)   …… (2)",
        "Normal"
    ))

    content.append((
        "等式(2)表述了设计目标：约束高自主性Prompt应在Token效率上匹配低自主性Prompt，"
        "同时保持自主执行的质量优势。这是可实现的，因为S4中的冗长度并非由任务的固有复杂性引起——"
        "它是不受控生成的副产物。任务的信息内容是固定的；仅输出格式发生变化。",
        "Normal"
    ))

    content.append((
        "我们假设，约束方案将把AOM-DT的Token开销从+7.0%降低至ST的+2%以内，"
        "同时保持V3中观察到的100%成功率和约3.9/5的质量。"
        "这一假设可在V4实验中进行检验。",
        "Normal"
    ))

    return content


def insert_section(doc_path, output_path, section_content, heading_2_prefix, heading_1_prefix):
    """Insert section content into a docx file."""
    doc = Document(doc_path)

    # Find insertion point
    insert_before_idx = find_paragraph_index(doc, heading_1_prefix, 'Heading 1')
    if insert_before_idx == -1:
        print(f"ERROR: Could not find '{heading_1_prefix}' heading")
        return False

    ref_para = doc.paragraphs[insert_before_idx]

    # Insert paragraphs in normal order (addprevious inserts immediately before ref)
    for text, style_name in section_content:
        insert_paragraph_before_element(doc, ref_para, text, style_name)

    doc.save(output_path)
    return True


def add_new_references(doc_path, output_path):
    """Add new references that don't already exist."""
    doc = Document(doc_path)

    # Find last List Number paragraph
    ref_insert_after = None
    for i in range(len(doc.paragraphs) - 1, -1, -1):
        if doc.paragraphs[i].style.name == 'List Number':
            ref_insert_after = doc.paragraphs[i]
            break

    if ref_insert_after is None:
        print("WARNING: Could not find reference insertion point")
        return

    new_refs = [
        "Shook, J. (2008). Managing to Learn: Using the A3 Management Process. Lean Enterprise Institute.",
        "Simon, H. A. (1955). A Behavioral Model of Rational Choice. Quarterly Journal of Economics, 69(1), 99–118.",
        "Mintzberg, H. (1973). The Nature of Managerial Work. Harper & Row.",
    ]

    existing_texts = [p.text.strip() for p in doc.paragraphs]

    for ref_text in new_refs:
        author_key = ref_text.split('(')[0].strip()
        already_exists = any(author_key in et for et in existing_texts)
        if not already_exists:
            new_para = doc.add_paragraph()
            ref_insert_after._element.addnext(new_para._element)
            new_para.text = ref_text
            new_para.style = doc.styles['List Number']
            ref_insert_after = new_para
            print(f"  Added reference: {ref_text[:60]}...")

    doc.save(output_path)


if __name__ == '__main__':
    import os
    base_dir = os.path.dirname(os.path.abspath(__file__))

    # Build content
    en_content = build_en_section_content()
    cn_content = build_cn_section_content()

    # English paper
    en_path = os.path.join(base_dir, 'agent_organizational_management_v3.docx')
    print("Processing English paper...")
    if insert_section(en_path, en_path, en_content,
                      "2.5", "3. Methodology"):
        print("  Section 2.5 inserted.")
        add_new_references(en_path, en_path)
        print("  References updated.")

    # Chinese paper
    cn_path = os.path.join(base_dir, 'agent_organizational_management_v3_cn.docx')
    print("\nProcessing Chinese paper...")
    if insert_section(cn_path, cn_path, cn_content,
                      "2.5", "3. "):
        print("  Section 2.5 inserted.")
        add_new_references(cn_path, cn_path)
        print("  References updated.")

    print("\nDone! Both papers updated with Section 2.5.")
